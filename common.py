from langchain_community.document_loaders import YoutubeLoader,WebBaseLoader
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
load_dotenv()
import comtypes.client
import streamlit as st
from langchain_core.prompts import ChatMessagePromptTemplate
from IPython.display import Markdown, display
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
import os
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import markdown
from nltk.tokenize import word_tokenize
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.summarize import load_summarize_chain
from langchain.chains.question_answering import load_qa_chain
from langchain_core.runnables import RunnableParallel
import base64
from langchain.docstore.document import Document
models=ChatGoogleGenerativeAI(model='gemini-1.5-pro')
languages = {
    "ab": "Abkhazian",
    "aa": "Afar",
    "af": "Afrikaans",
    "ak": "Akan",
    "sq": "Albanian",
    "am": "Amharic",
    "ar": "Arabic",
    "hy": "Armenian",
    "as": "Assamese",
    "ay": "Aymara",
    "az": "Azerbaijani",
    "bn": "Bangla",
    "ba": "Bashkir",
    "eu": "Basque",
    "be": "Belarusian",
    "bho": "Bhojpuri",
    "bs": "Bosnian",
    "br": "Breton",
    "bg": "Bulgarian",
    "my": "Burmese",
    "ca": "Catalan",
    "ceb": "Cebuano",
    "zh-Hans": "Chinese (Simplified)",
    "zh-Hant": "Chinese (Traditional)",
    "co": "Corsican",
    "hr": "Croatian",
    "cs": "Czech",
    "da": "Danish",
    "dv": "Divehi",
    "nl": "Dutch",
    "dz": "Dzongkha",
    "en": "English",
     'en-US' : 'English (United States)',
     'en-IN':'English(Indian)',
    "eo": "Esperanto",
    "et": "Estonian",
    "ee": "Ewe",
    "fo": "Faroese",
    "fj": "Fijian",
    "fil": "Filipino",
    "fi": "Finnish",
    "fr": "French",
    "gaa": "Ga",
    "gl": "Galician",
    "lg": "Ganda",
    "ka": "Georgian",
    "de": "German",
    "el": "Greek",
    "gn": "Guarani",
    "gu": "Gujarati",
    "ht": "Haitian Creole",
    "ha": "Hausa",
    "haw": "Hawaiian",
    "iw": "Hebrew",
    "hi": "Hindi",
    "hmn": "Hmong",
    "hu": "Hungarian",
    "is": "Icelandic",
    "ig": "Igbo",
    "id": "Indonesian",
    "ga": "Irish",
    "it": "Italian",
    "ja": "Japanese",
    "jv": "Javanese",
    "kl": "Kalaallisut",
    "kn": "Kannada",
    "kk": "Kazakh",
    "kha": "Khasi",
    "km": "Khmer",
    "rw": "Kinyarwanda",
    "ko": "Korean",
    "kri": "Krio",
    "ku": "Kurdish",
    "ky": "Kyrgyz",
    "lo": "Lao",
    "la": "Latin",
    "lv": "Latvian",
    "ln": "Lingala",
    "lt": "Lithuanian",
    "luo": "Luo",
    "lb": "Luxembourgish",
    "mk": "Macedonian",
    "mg": "Malagasy",
    "ms": "Malay",
    "ml": "Malayalam",
    "mt": "Maltese",
    "gv": "Manx",
    "mi": "Māori",
    "mr": "Marathi",
    "mn": "Mongolian",
    "mfe": "Morisyen",
    "ne": "Nepali",
    "new": "Newari",
    "nso": "Northern Sotho",
    "no": "Norwegian",
    "ny": "Nyanja",
    "oc": "Occitan",
    "or": "Odia",
    "om": "Oromo",
    "os": "Ossetic",
    "pam": "Pampanga",
    "ps": "Pashto",
    "fa": "Persian",
    "pl": "Polish",
    "pt": "Portuguese",
    "pt-PT": "Portuguese (Portugal)",
    "pa": "Punjabi",
    "qu": "Quechua",
    "ro": "Romanian",
    "rn": "Rundi",
    "ru": "Russian",
    "sm": "Samoan",
    "sg": "Sango",
    "sa": "Sanskrit",
    "gd": "Scottish Gaelic",
    "sr": "Serbian",
    "crs": "Seselwa Creole French",
    "sn": "Shona",
    "sd": "Sindhi",
    "si": "Sinhala",
    "sk": "Slovak",
    "sl": "Slovenian",
    "so": "Somali",
    "st": "Southern Sotho",
    "es": "Spanish",
    "su": "Sundanese",
    "sw": "Swahili",
    "ss": "Swati",
    "sv": "Swedish",
    "tg": "Tajik",
    "ta": "Tamil",
    "tt": "Tatar",
    "te": "Telugu",
    "th": "Thai",
    "bo": "Tibetan",
    "ti": "Tigrinya",
    "to": "Tongan",
    "ts": "Tsonga",
    "tn": "Tswana",
    "tum": "Tumbuka",
    "tr": "Turkish",
    "tk": "Turkmen",
    "uk": "Ukrainian",
    "ur": "Urdu",
    "ug": "Uyghur",
    "uz": "Uzbek",
    "ve": "Venda",
    "vi": "Vietnamese",
    "war": "Waray",
    "cy": "Welsh",
    "fy": "Western Frisian",
    "wo": "Wolof",
    "xh": "Xhosa",
    "yi": "Yiddish",
    "yo": "Yoruba",
    "zu": "Zulu"
}

def html_content_parser(output):
    markdon_to_html=markdown.markdown(output)
    soup = BeautifulSoup(markdon_to_html, "html.parser")
    return soup

def markdown_result(result):
    return display(Markdown(result))
def tokenize(lang):
    return word_tokenize(lang)
from langchain.document_loaders import YoutubeLoader
def youtubetranscript(links,language=None,translation='en'):
    
    text=''
    if language=="en":
    
        for i in links.split(','):
            loader=YoutubeLoader.from_youtube_url(i,add_video_info=True)
            Transcript=loader.load()
            text+=Transcript[0].page_content
        return text
    else:
        for i in links.split(','):
            loader=YoutubeLoader.from_youtube_url(i,language=language,translation=translation,add_video_info=True)
            Transcript=loader.load()
            text+=Transcript[0].page_content
        return text



    

def videogenre(models,transcript):
    
    prompt = '''
    Based on the provided YouTube video transcript, classify the genre into one of the following categories:
    - Case Study
    - Educational/Tutorial
    - Documentary
    - Podcast/Interview/Debate
    - News/Current Affairs
    - Other

    Please provide only the genre from the list above.

    Input Transcript:
    {Transcript}
    '''
    template = PromptTemplate(input_variables=['Transcript'], template=prompt)
    content_genre=models.invoke(template.format(Transcript=transcript)).content
    return content_genre
def web_data(link):
    text=''
    for i in link.split(','):
        loader=WebBaseLoader(i)
        text+=loader.load()[0].page_content
    return text
def add_html_to_docx(html, doc):
    for element in html:
        if element.name == "h2":
            run = doc.add_heading(level=1).add_run(element.text)
            run.font.size = Pt(24)
        elif element.name == "p":
            p = doc.add_paragraph()
            run = p.add_run(element.text)
        elif element.name == "ul":
            for li in element.find_all("li"):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(li.text)
def split_text(text):
    text_split=RecursiveCharacterTextSplitter(chunk_size=4000,chunk_overlap=500)
    return text_split.split_text(text)

def qna_chain(models,text,question):
    documents = [Document(page_content=text)]
    
    chain=load_qa_chain(models,chain_type='map_reduce', verbose=True, return_intermediate_steps=False)
    result=chain({'input_documents':documents,'question':question},return_only_outputs=True)
    return result['output_text']

def report_creation(models, text, format, type_of_report):
    report_template = """You are tasked to write {format} for a report, use {format} as the respective heading. The report will be a {type_of_report} based on the following data:
    '{text}'
    """

    prompt = PromptTemplate(template=report_template, input_variables=['text', 'format', 'type_of_report'])
    
    format_components = format.split(',')
    formatted_prompts = []
    
    for component in format_components:
        formatted_prompt = prompt.format(format=component, text=text, type_of_report=type_of_report)
        formatted_prompts.append(models.invoke(formatted_prompt).content)  

    
    results = formatted_prompts  
    
    output = ''
    for result in results:
        output += result  # Assuming `result` is already the text output

    return output

def summarize_video(Transcript, models, content_genre):
    content_genre = content_genre.replace('\n', '').strip()
    token_len=models.get_num_tokens(Transcript)
    
    if content_genre == 'Educational/Tutorial':
        prompt = f'''Summarize this transcript in points and briefly describe it. If the video involves coding, include the piece of code as well. from the following:
        {Transcript} Replace the word Transcript with Video in result'''
    
    elif content_genre in ['Case Study', 'News/Current Affairs']:
        prompt = f''' Summarize this transcript in detail, describing each event in the order of occurrence.from the following:
        {Transcript} Replace the word Transcript with Video in result'''
        
    elif content_genre in ['Podcast/Interview/Debate', 'Documentary']:
        prompt = f'''Explain this transcript in detail with subheadings and points from the following :
        {Transcript} Replace the word Transcript with Video in result'''
        
    else:
        prompt = f''' Provide a general summary of this transcript. from the following:
        {Transcript} Replace the word Transcript with Video in result'''

    
    if token_len<5000:
        template = PromptTemplate(input_variables=['Transcript'], template=prompt)
        result = models.invoke(template.format(Transcript=Transcript)).content
        return result
    else:
        combine_prompt = """
    Summarize the following text, breaking it down into key sections with bullet points under each heading. Ensure each section highlights the most important details.
    
    TEXT:
    "{text}"
    
    SUMMARY FORMAT:
    
    - **Introduction:**
      - Key points of the introduction.
      
    - **Main Points:**
      - Key details about the main topic.
      
    - **Conclusion:**
      - Key points of the conclusion.
    """
        text_splitter = RecursiveCharacterTextSplitter(separators=["\n\n", "\n"], chunk_size=5000, chunk_overlap=500)
        summary_prompt_template = PromptTemplate(template=prompt, input_variables=['text'])
        combine_prompt_template = PromptTemplate(template=combine_prompt, input_variables=['text'])
        docs = text_splitter.create_documents([Transcript])
        summary_chain = load_summarize_chain(llm=models, chain_type='map_reduce', map_prompt=summary_prompt_template, combine_prompt=combine_prompt_template)
        
        # Running the summarization chain with the provided text
        output = summary_chain(docs)
        
        return output['output_text']




def summarize_web_and_vid_or_web(Text, models):
    token_len=models.get_num_tokens(Text)
    
    documents = [Document(page_content=Text)]
    prompt = """Write the concise summary of the following:
    "{text}" 
    CONCISE SUMMARY:
    """
    
    combine_prompt = """
    Summarize the following text, breaking it down into key sections with bullet points under each heading. Ensure each section highlights the most important details.
    
    TEXT:
    "{text}"
    
    SUMMARY FORMAT:
    
    - **Introduction:**
      - Key points of the introduction.
      
    - **Main Points:**
      - Key details about the main topic.
      
    - **Conclusion:**
      - Key points of the conclusion.
    """
    
    # Correctly creating PromptTemplate objects
    summary_prompt_template = PromptTemplate(template=prompt, input_variables=['text'])
    combine_prompt_template = PromptTemplate(template=combine_prompt, input_variables=['text'])
    
    # Loading the summarize chain
    if token_len<5000:
        summary_chain = load_summarize_chain(llm=models, chain_type='map_reduce', map_prompt=summary_prompt_template, combine_prompt=combine_prompt_template)
        
        # Running the summarization chain with the provided text
        output = summary_chain(documents)
        
        return output['output_text']
    else:
        text_splitter = RecursiveCharacterTextSplitter(separators=["\n\n", "\n"], chunk_size=5000, chunk_overlap=500)

        docs = text_splitter.create_documents([Text])
        summary_chain = load_summarize_chain(llm=models, chain_type='map_reduce', map_prompt=summary_prompt_template, combine_prompt=combine_prompt_template)
        
        # Running the summarization chain with the provided text
        output = summary_chain(docs)
        
        return output['output_text']

def match_back_language(value):
    for key,values in languages.items():
        if values==value:
            return key
def doc_to_pdf(word_path='report.docx',pdf_path='report.pdf'):
    word_path='report.docx'
    pdf_path='report.pdf'
    doc=Document(word_path)
    word=comtypes.client.CreateObject('Word.Application')
    docx_path=os.path.abspath(word_path)
    pdf_path=os.path.abspath(pdf_path)
    pdf_format=17
    word.Visible=False
    in_file=word.Documents.Open(docx_path)
    in_file.SaveAs(pdf_path,File_Format=pdf_format)
    in_file.Close()
    word.Quit
    os.remove(word_path)

def displayPDF(file):
    # Opening file from file path
    with open(file, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')

    # Embedding PDF in HTML
    pdf_display = F'<embed src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf">'

    # Displaying File
    st.markdown(pdf_display, unsafe_allow_html=True)




